import os
import PyPDF2
from docx import Document as DocxDocument
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.chains import RetrievalQA
from langchain.llms import HuggingFaceHub
from dotenv import load_dotenv

load_dotenv()
huggingAPI = os.getenv('HUGGING_FACE_API')

def extractPDF(file_path):
    text = ""
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        for page in range(len(reader.pages)):
            text += reader.pages[page].extract_text()
    return text

def extractWord(file_path):
    doc = DocxDocument(file_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

def processDocument(file_path):
    _, ext = os.path.splitext(file_path)
    if ext == '.pdf':
        return extractPDF(file_path)
    elif ext == '.docx':
        return extractWord(file_path)
    else:
        raise ValueError("Unsupported file format. Please upload a .pdf or .docx file.")

def createChunks(text):
    splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=0)
    documents = splitter.split_text(text)
    return [Document(page_content=doc) for doc in documents]

def createVector(documents):
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")
    vector_store = FAISS.from_documents(documents, embeddings)
    return vector_store.as_retriever()

def setupRAG(vector_store):
    llm = HuggingFaceHub(
        repo_id="google/flan-t5-large",huggingfacehub_api_token=huggingAPI
    )
    qa_chain = RetrievalQA.from_chain_type(llm=llm, retriever=vector_store, chain_type="stuff")
    return qa_chain

def RAGBot(file_path, user_query):
    document_text = processDocument(file_path)
    document_chunks = createChunks(document_text)
    vector_store = createVector(document_chunks)
    qa_chain = setupRAG(vector_store)
    result = qa_chain({"query": user_query})
    return result['result']